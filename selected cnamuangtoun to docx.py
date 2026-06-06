"""
Resume Text → DOCX Converter (LLM-based)
==========================================
Reads : selected_dataset.csv
Writes: cv_files/ folder  +  final_dataset.csv

Uses GPT-4o-mini to parse raw resume text into clean JSON,
then renders a professionally styled DOCX from that JSON.
"""

import os
import re
import json
import random
import pandas as pd
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

random.seed(42)

# ── CONFIG ─────────────────────────────────────────────────────
OPENAI_API_KEY = " "           # ← paste your key
MODEL          = "gpt-4o-mini"
INPUT_FILE     = "selected_dataset.csv"
OUTPUT_DIR     = "cv_files"
OUTPUT_CSV     = "final_dataset.csv"

os.makedirs(OUTPUT_DIR, exist_ok=True)
client = OpenAI(api_key=OPENAI_API_KEY)

# ── STYLE POOLS ────────────────────────────────────────────────
ACCENT_COLORS = [
    RGBColor(0x1a, 0x3c, 0x5e),   # Navy
    RGBColor(0x2d, 0x5a, 0x3f),   # Forest green
    RGBColor(0x6b, 0x2e, 0x3e),   # Burgundy
    RGBColor(0x2c, 0x6e, 0x6e),   # Teal
    RGBColor(0x4a, 0x2c, 0x6e),   # Purple
    RGBColor(0x8b, 0x45, 0x13),   # Brown
    RGBColor(0x1e, 0x5f, 0x6e),   # Deep cyan
    RGBColor(0x3a, 0x6b, 0x4b),   # Sage green
]
GREY       = RGBColor(0x55, 0x55, 0x55)
DARK_GREY  = RGBColor(0x22, 0x22, 0x22)
SEPARATORS = [' | ', ' • ', ' — ', '  ·  ']
MARGINS    = [0.75, 0.80, 0.90, 1.0]

# ── LLM PARSING ────────────────────────────────────────────────
PARSE_PROMPT = """You are a resume parser. Parse the following raw resume text into clean structured JSON.

The text may have section headers glued to content (e.g. "SummaryJohn is a..."),
repeated section names (e.g. "Experience" appearing once per job entry),
and inconsistent formatting. Fix all of these.

Return ONLY valid JSON in exactly this structure — no text before or after:
{
  "name": "Full Name",
  "email": "email or empty string",
  "phone": "phone or empty string",
  "location": "city, state or empty string",
  "summary": "2-4 sentence professional summary paragraph",
  "skills": ["skill1", "skill2", "skill3"],
  "experience": [
    {
      "title": "Job Title",
      "company": "Company Name",
      "location": "City, State",
      "dates": "MM/YYYY - MM/YYYY or Present",
      "bullets": [
        "Achievement or responsibility one",
        "Achievement or responsibility two"
      ]
    }
  ],
  "education": [
    {
      "degree": "Degree Name and Field",
      "institution": "University Name",
      "year": "Graduation year or expected"
    }
  ],
  "certifications": ["cert1", "cert2"],
  "languages": ["Language (level)"]
}

Rules:
- name: extract the person's full name from the top of the resume
- summary: write ONE clean paragraph — do NOT use the section header word "Summary"
- skills: flat list of individual skills — no categories, no bullets
- experience: group all entries under one list — do NOT repeat per job
- bullets: minimum 3 per job entry, clean sentences without bullet symbols
- certifications and languages: empty list [] if not present
- If a field cannot be found, use empty string "" or empty list []

RAW RESUME TEXT:
"""

def parse_resume_with_llm(resume_text):
    """Send raw resume text to GPT-4o-mini and get clean JSON back."""
    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {"role": "user", "content": PARSE_PROMPT + resume_text}
        ],
        temperature=0.1,
    )
    raw = response.choices[0].message.content.strip()
    # Strip markdown fences if present
    raw = re.sub(r"```(?:json)?|```", "", raw).strip()
    # Find JSON object
    start = raw.find("{")
    end   = raw.rfind("}") + 1
    if start != -1 and end > start:
        raw = raw[start:end]
    return json.loads(raw)


# ── DOCX UTILITIES ─────────────────────────────────────────────
def add_divider(paragraph, accent, thickness="6"):
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    thickness)
    bot.set(qn("w:color"),
            f"{accent[0]:02x}{accent[1]:02x}{accent[2]:02x}")
    pBdr.append(bot)
    pPr.append(pBdr)

def sp(p, before=0, after=3):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)


# ── DOCX WRITER ────────────────────────────────────────────────
def write_docx(cv, filepath):
    # Random style choices per CV
    accent     = random.choice(ACCENT_COLORS)
    sep        = random.choice(SEPARATORS)
    margin     = random.choice(MARGINS)
    name_size  = random.choice([16, 18, 20])
    body_size  = random.choice([9.5, 10.0])
    sec_size   = random.choice([10, 11])
    center_hdr = random.choice([True, True, False])
    div_thick  = random.choice(["6", "8"])

    doc = Document()
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Inches(margin)
        sec.top_margin  = sec.bottom_margin = Inches(0.75)

    # ── NAME ─────────────────────────────────────────────────
    name = cv.get("name", "Candidate")
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(name_size)
    r.font.color.rgb = accent
    p.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if center_hdr
        else WD_ALIGN_PARAGRAPH.LEFT
    )
    sp(p, before=0, after=3)

    # ── CONTACT LINE ─────────────────────────────────────────
    contact_parts = [
        cv.get("email", ""),
        cv.get("phone", ""),
        cv.get("location", ""),
    ]
    contact_parts = [c for c in contact_parts if c.strip()]
    if contact_parts:
        p = doc.add_paragraph()
        r = p.add_run(sep.join(contact_parts))
        r.font.size = Pt(9)
        r.font.color.rgb = GREY
        p.paragraph_format.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if center_hdr
            else WD_ALIGN_PARAGRAPH.LEFT
        )
        sp(p, after=8)
        add_divider(p, accent, div_thick)

    def section_heading(label):
        p = doc.add_paragraph()
        r = p.add_run(label.upper())
        r.bold = True
        r.font.size = Pt(sec_size)
        r.font.color.rgb = accent
        sp(p, before=10, after=3)
        add_divider(p, accent, div_thick)

    def body_para(text, italic=False, bold=False, color=None):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(body_size)
        r.italic = italic
        r.bold   = bold
        if color:
            r.font.color.rgb = color
        sp(p, after=2)
        return p

    def bullet_para(text):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.size = Pt(body_size)
        sp(p, after=1)

    # ── SUMMARY ──────────────────────────────────────────────
    summary = cv.get("summary", "")
    if summary:
        section_heading("Professional Summary")
        body_para(summary)

    # ── SKILLS ───────────────────────────────────────────────
    skills = cv.get("skills", [])
    if skills:
        section_heading("Skills")
        # Group skills into rows of 4
        rows = [skills[i:i+4] for i in range(0, len(skills), 4)]
        for row in rows:
            body_para("  •  ".join(row))

    # ── EXPERIENCE ───────────────────────────────────────────
    experience = cv.get("experience", [])
    if experience:
        section_heading("Work Experience")
        for job in experience:
            title   = job.get("title", "")
            company = job.get("company", "")
            loc     = job.get("location", "")
            dates   = job.get("dates", "")
            bullets = job.get("bullets", [])

            # Job title
            p = doc.add_paragraph()
            r = p.add_run(title)
            r.bold = True
            r.font.size = Pt(body_size + 0.5)
            r.font.color.rgb = DARK_GREY
            sp(p, before=6, after=1)

            # Company | Location | Dates
            meta_parts = [c for c in [company, loc, dates] if c.strip()]
            if meta_parts:
                p = doc.add_paragraph()
                r = p.add_run("  |  ".join(meta_parts))
                r.font.size = Pt(body_size - 0.5)
                r.font.color.rgb = GREY
                r.italic = True
                sp(p, after=3)

            # Bullets
            for b in bullets:
                b_clean = b.strip().lstrip("-•*·► ")
                if b_clean:
                    bullet_para(b_clean)

            # Small gap after each job
            sp(doc.add_paragraph(), after=2)

    # ── EDUCATION ────────────────────────────────────────────
    education = cv.get("education", [])
    if education:
        section_heading("Education")
        for edu in education:
            degree  = edu.get("degree", "")
            inst    = edu.get("institution", "")
            year    = edu.get("year", "")

            if degree:
                p = doc.add_paragraph()
                r = p.add_run(degree)
                r.bold = True
                r.font.size = Pt(body_size)
                r.font.color.rgb = DARK_GREY
                sp(p, before=4, after=1)

            meta = [c for c in [inst, year] if c.strip()]
            if meta:
                p = doc.add_paragraph()
                r = p.add_run("  |  ".join(meta))
                r.font.size = Pt(body_size - 0.5)
                r.font.color.rgb = GREY
                r.italic = True
                sp(p, after=3)

    # ── CERTIFICATIONS ───────────────────────────────────────
    certs = cv.get("certifications", [])
    if certs:
        section_heading("Certifications")
        for c in certs:
            if c.strip():
                bullet_para(c.strip())

    # ── LANGUAGES ────────────────────────────────────────────
    langs = cv.get("languages", [])
    if langs:
        section_heading("Languages")
        body_para("  •  ".join(langs))

    doc.save(filepath)


# ── MAIN ───────────────────────────────────────────────────────
def main():
    df = pd.read_csv(INPUT_FILE)

    # Add resume_id if not present
    if "resume_id" not in df.columns:
        label_code = {"Good Fit": "S", "Potential Fit": "P", "No Fit": "W"}
        counters   = {}
        ids        = []
        for _, row in df.iterrows():
            key = (row["jd_id"], row["label"])
            counters[key] = counters.get(key, 0) + 1
            code = label_code.get(row["label"], "X")
            ids.append(f"{row['jd_id']}_{code}{counters[key]:02d}")
        df["resume_id"] = ids

    print(f"Loaded {len(df)} resumes")
    print(f"Model : {MODEL}")
    print(f"Output: {OUTPUT_DIR}/")
    print()

    file_paths = []
    failed     = []

    for i, row in df.iterrows():
        rid      = row["resume_id"]
        filepath = os.path.join(OUTPUT_DIR, f"{rid}.docx")

        try:
            # Step 1 — LLM parses raw text → clean JSON
            cv_json = parse_resume_with_llm(row["resume_text"])

            # Step 2 — Write DOCX from JSON
            write_docx(cv_json, filepath)

            file_paths.append(filepath)
            print(f"  [{i+1:3}/{len(df)}] ✅  {rid}  — {cv_json.get('name','?')}")

        except Exception as e:
            print(f"  [{i+1:3}/{len(df)}] ❌  {rid}  — {e}")
            file_paths.append(None)
            failed.append(rid)

    df["file_path"] = file_paths

    # Save final dataset
    cols = [c for c in ["resume_id","jd_id","label","label_num",
                         "file_path","resume_text","job_description_text"]
            if c in df.columns]
    df[cols].to_csv(OUTPUT_CSV, index=False)

    print()
    print("=" * 50)
    print(f"Done.")
    print(f"  Generated : {len(df) - len(failed)} DOCX files")
    print(f"  Failed    : {len(failed)}")
    print(f"  Saved     : {OUTPUT_CSV}")
    if failed:
        print(f"  Failed IDs: {failed}")


if __name__ == "__main__":
    main()
